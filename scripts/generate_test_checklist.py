#!/usr/bin/env python3
"""
Generate a comprehensive test checklist Word document for Asset Management App
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime

def create_test_checklist():
    doc = Document()
    
    # Title
    title = doc.add_heading('Asset Management App - Test Checklist', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.italic = True
    subtitle_format.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Table of Contents section
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Authentication & User Management',
        '2. Dashboard',
        '3. Search & Inventory',
        '4. Asset Management',
        '5. Asset Types',
        '6. Certificates/Documents',
        '7. Activity Log',
        '8. QR Code Features',
        '9. Quick Actions & Shortcuts',
        '10. Admin Features',
        '11. Profile & Settings',
        '12. Mobile-Specific Features',
        '13. Web-Specific Features',
        '14. Performance & Error Handling'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== SECTION 1: Authentication ==========
    doc.add_heading('1. Authentication & User Management', 1)
    
    auth_table = doc.add_table(rows=1, cols=4)
    auth_table.style = 'Light Grid Accent 1'
    auth_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = auth_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    # Format header
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    auth_tests = [
        ('Login Screen', '1. Open app\n2. Enter valid email/password\n3. Click Login\n4. Verify redirect to dashboard', '', ''),
        ('Login - Invalid Credentials', '1. Enter wrong email/password\n2. Verify error message\n3. Verify no redirect', '', ''),
        ('Login - Empty Fields', '1. Leave fields empty\n2. Click Login\n3. Verify validation message', '', ''),
        ('Registration', '1. Navigate to Register\n2. Fill all required fields\n3. Submit\n4. Verify account creation', '', ''),
        ('Forgot Password', '1. Click "Forgot Password"\n2. Enter email\n3. Verify reset email sent', '', ''),
        ('Logout', '1. Click Logout button\n2. Verify redirect to login\n3. Verify session cleared', '', ''),
        ('Session Persistence', '1. Login\n2. Close app\n3. Reopen app\n4. Verify still logged in', '', ''),
        ('Auto-logout on Token Expiry', '1. Login\n2. Wait for token expiry\n3. Perform action\n4. Verify redirect to login', '', ''),
    ]
    
    for feature, steps, working, feedback in auth_tests:
        row_cells = auth_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()  # Spacing
    
    # ========== SECTION 2: Dashboard ==========
    doc.add_heading('2. Dashboard', 1)
    
    dashboard_table = doc.add_table(rows=1, cols=4)
    dashboard_table.style = 'Light Grid Accent 1'
    
    hdr_cells = dashboard_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    dashboard_tests = [
        ('Dashboard Load', '1. Login\n2. Verify dashboard displays\n3. Check all sections visible', '', ''),
        ('My Tasks Section', '1. View tasks list\n2. Scroll through tasks\n3. Verify task details display', '', ''),
        ('Task Actions', '1. Click on a task\n2. Verify action modal opens\n3. Complete task\n4. Verify task removed from list', '', ''),
        ('Recent Assets', '1. Check recent assets section\n2. Verify asset cards display\n3. Click asset\n4. Verify navigation to asset detail', '', ''),
        ('Quick Actions - Search', '1. Click Search button\n2. Verify navigation to search screen', '', ''),
        ('Quick Actions - Certs', '1. Click Certs button\n2. Verify navigation to certs screen', '', ''),
        ('Quick Actions - Activity', '1. Click Activity button\n2. Verify navigation to activity screen', '', ''),
        ('Shortcuts Section', '1. View shortcuts grid\n2. Verify custom shortcuts display\n3. Click shortcut\n4. Verify action executes', '', ''),
        ('Add Shortcut', '1. Click "Add Shortcut"\n2. Select shortcut type\n3. Verify shortcut added\n4. Verify appears in grid', '', ''),
        ('Remove Shortcut', '1. Click "Manage Added"\n2. Remove a shortcut\n3. Verify removed from grid', '', ''),
        ('Dashboard Navigation (Web)', '1. Click navbar links\n2. Verify navigation works\n3. Verify active state highlighting', '', ''),
    ]
    
    for feature, steps, working, feedback in dashboard_tests:
        row_cells = dashboard_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 3: Search & Inventory ==========
    doc.add_heading('3. Search & Inventory', 1)
    
    search_table = doc.add_table(rows=1, cols=4)
    search_table.style = 'Light Grid Accent 1'
    
    hdr_cells = search_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    search_tests = [
        ('Search Input', '1. Enter search query\n2. Verify results filter\n3. Verify real-time search', '', ''),
        ('Quick Filters', '1. Click "My Assets"\n2. Verify filtered results\n3. Click "Needs Service"\n4. Verify filtered results\n5. Click "In Service"\n6. Verify filtered results', '', ''),
        ('QR Awaiting Filter', '1. Click "QR Awaiting"\n2. Verify only UUID assets shown\n3. Verify "QR awaiting" label', '', ''),
        ('Advanced Filters', '1. Open Filters modal\n2. Select asset types\n3. Select status\n4. Apply filters\n5. Verify results', '', ''),
        ('Clear Filters', '1. Apply filters\n2. Click "Clear All"\n3. Verify all filters reset', '', ''),
        ('Sort Options', '1. Select sort option\n2. Verify results sorted\n3. Change sort order\n4. Verify re-sorted', '', ''),
        ('Grid View', '1. Click Grid view\n2. Verify card layout\n3. Verify cards display correctly', '', ''),
        ('Table View', '1. Click Table view\n2. Verify table layout\n3. Verify all columns visible\n4. Verify horizontal scroll works', '', ''),
        ('Pagination', '1. Navigate through pages\n2. Verify page numbers\n3. Click "View All"\n4. Verify all results shown', '', ''),
        ('Asset ID Click', '1. Click Asset ID in table\n2. Verify navigation to asset detail', '', ''),
        ('Dynamic Columns', '1. Filter by asset type with custom fields\n2. Verify dynamic columns appear\n3. Verify values display correctly', '', ''),
        ('Mobile Search Layout', '1. Open on mobile\n2. Verify responsive layout\n3. Verify cards display properly\n4. Verify filters accessible', '', ''),
    ]
    
    for feature, steps, working, feedback in search_tests:
        row_cells = search_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 4: Asset Management ==========
    doc.add_heading('4. Asset Management', 1)
    
    asset_table = doc.add_table(rows=1, cols=4)
    asset_table.style = 'Light Grid Accent 1'
    
    hdr_cells = asset_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    asset_tests = [
        ('Asset Detail View', '1. Navigate to asset\n2. Verify all fields display\n3. Verify images load\n4. Verify documents list', '', ''),
        ('Edit Asset', '1. Click Edit\n2. Modify fields\n3. Save\n4. Verify changes saved\n5. Verify updated in list', '', ''),
        ('Create New Asset', '1. Click "Create New Asset"\n2. Fill required fields\n3. Add optional fields\n4. Submit\n5. Verify asset created', '', ''),
        ('Asset Status Change', '1. Change asset status\n2. Verify status updates\n3. Verify activity logged', '', ''),
        ('Assign Asset', '1. Assign asset to user\n2. Verify assignment saved\n3. Verify appears in user\'s assets', '', ''),
        ('Asset Images', '1. Upload image\n2. Verify image displays\n3. Verify thumbnail in list', '', ''),
        ('Asset Documents', '1. Upload document\n2. Verify document appears\n3. Click "View"\n4. Verify document opens', '', ''),
        ('Additional Fields', '1. View additional fields\n2. Verify custom fields display\n3. Edit custom field values\n4. Verify saved', '', ''),
        ('Document History', '1. View document history\n2. Verify all documents listed\n3. Verify dates correct\n4. Verify "View" links work', '', ''),
        ('Service/Repair Reports', '1. Sign off service/repair\n2. Attach report\n3. Verify appears in certs\n4. Verify "Not provided" if missing', '', ''),
        ('Attach Report Later', '1. View asset with missing report\n2. Click "Attach report"\n3. Upload document\n4. Verify appears in certs', '', ''),
        ('Back Navigation', '1. Navigate to asset from search\n2. Click back\n3. Verify returns to search\n4. Test from different screens', '', ''),
    ]
    
    for feature, steps, working, feedback in asset_tests:
        row_cells = asset_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 5: Asset Types ==========
    doc.add_heading('5. Asset Types', 1)
    
    type_table = doc.add_table(rows=1, cols=4)
    type_table.style = 'Light Grid Accent 1'
    
    hdr_cells = type_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    type_tests = [
        ('Asset Type List', '1. Navigate to Inventory\n2. Click Asset Types tab\n3. Verify types listed\n4. Verify images display', '', ''),
        ('Asset Type Detail', '1. Click asset type\n2. Verify details display\n3. Verify custom fields listed', '', ''),
        ('Create Asset Type', '1. Click "Create New Asset Type"\n2. Fill form\n3. Add custom fields\n4. Submit\n5. Verify created', '', ''),
        ('Edit Asset Type', '1. Click Edit\n2. Modify fields\n3. Add/remove custom fields\n4. Save\n5. Verify changes', '', ''),
        ('Custom Fields', '1. Add custom field\n2. Set field type\n3. Set required/optional\n4. Verify appears in asset forms', '', ''),
        ('Asset Type Image', '1. Upload image for type\n2. Verify displays in list\n3. Verify displays in asset forms', '', ''),
    ]
    
    for feature, steps, working, feedback in type_tests:
        row_cells = type_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 6: Certificates/Documents ==========
    doc.add_heading('6. Certificates/Documents', 1)
    
    certs_table = doc.add_table(rows=1, cols=4)
    certs_table.style = 'Light Grid Accent 1'
    
    hdr_cells = certs_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    certs_tests = [
        ('Certs List View', '1. Navigate to Certs\n2. Verify documents listed\n3. Verify card/table layout', '', ''),
        ('Quick Filters - My Documents', '1. Click "My Documents"\n2. Verify only user\'s assets shown\n3. Verify correct matching', '', ''),
        ('Quick Filters - Expiring Soon', '1. Click "Expiring Soon"\n2. Verify filtered results\n3. Verify dates correct', '', ''),
        ('Quick Filters - Expired', '1. Click "Expired"\n2. Verify expired documents shown', '', ''),
        ('Document Filters', '1. Open filters\n2. Select document type\n3. Apply\n4. Verify filtered', '', ''),
        ('Open Document', '1. Click "Open" on document\n2. Verify document opens\n3. Verify correct document', '', ''),
        ('Edit Document', '1. Click "Edit"\n2. Modify details\n3. Save\n4. Verify changes', '', ''),
        ('Document Sorting', '1. Change sort option\n2. Verify documents re-sorted', '', ''),
        ('Mobile Certs Layout', '1. Open on mobile\n2. Verify card layout\n3. Verify filters accessible\n4. Verify alignment correct', '', ''),
        ('Service/Repair Reports', '1. Verify service reports appear\n2. Verify repair reports appear\n3. Verify missing reports show "Not provided"', '', ''),
        ('Document Links', '1. Click document link\n2. Verify opens correctly\n3. Verify "View" text displays', '', ''),
    ]
    
    for feature, steps, working, feedback in certs_tests:
        row_cells = certs_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 7: Activity Log ==========
    doc.add_heading('7. Activity Log', 1)
    
    activity_table = doc.add_table(rows=1, cols=4)
    activity_table.style = 'Light Grid Accent 1'
    
    hdr_cells = activity_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    activity_tests = [
        ('Activity List', '1. Navigate to Activity\n2. Verify activities listed\n3. Verify chronological order', '', ''),
        ('Activity Details', '1. View activity entry\n2. Verify all details shown\n3. Verify user info\n4. Verify asset info', '', ''),
        ('Activity Filtering', '1. Filter by activity type\n2. Verify filtered results\n3. Filter by user\n4. Verify filtered', '', ''),
        ('Activity Navigation', '1. Click asset link\n2. Verify navigates to asset\n3. Verify back navigation works', '', ''),
    ]
    
    for feature, steps, working, feedback in activity_tests:
        row_cells = activity_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 8: QR Code Features ==========
    doc.add_heading('8. QR Code Features', 1)
    
    qr_table = doc.add_table(rows=1, cols=4)
    qr_table.style = 'Light Grid Accent 1'
    
    hdr_cells = qr_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    qr_tests = [
        ('QR Scanner', '1. Open QR scanner\n2. Grant camera permission\n3. Scan QR code\n4. Verify asset detected', '', ''),
        ('QR Code Display', '1. View asset\n2. Verify QR code displays\n3. Verify correct data encoded', '', ''),
        ('QR Sheet Generation', '1. Admin: Generate QR sheet\n2. Verify PDF generated\n3. Verify QR codes correct\n4. Verify printing works', '', ''),
        ('QR Check-in', '1. Scan QR code\n2. Verify check-in page opens\n3. Complete check-in\n4. Verify activity logged', '', ''),
        ('Location Capture', '1. Scan QR with location permission\n2. Verify location captured\n3. Verify saved to asset', '', ''),
    ]
    
    for feature, steps, working, feedback in qr_tests:
        row_cells = qr_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 9: Quick Actions & Shortcuts ==========
    doc.add_heading('9. Quick Actions & Shortcuts', 1)
    
    shortcuts_table = doc.add_table(rows=1, cols=4)
    shortcuts_table.style = 'Light Grid Accent 1'
    
    hdr_cells = shortcuts_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    shortcuts_tests = [
        ('Quick View', '1. Select Quick View shortcut\n2. Scan asset\n3. Verify asset detail opens\n4. Verify stays open', '', ''),
        ('Quick Transfer', '1. Select Quick Transfer\n2. Scan asset\n3. Select recipient\n4. Verify transfer completes\n5. Verify location captured', '', ''),
        ('Quick Transfer Office', '1. Select Quick Transfer Office\n2. Scan asset\n3. Verify assigned to admin\n4. Verify location captured\n5. Test already assigned message', '', ''),
        ('Transfer-To Me', '1. Select Transfer-To Me\n2. Scan asset\n3. Verify assigned to user\n4. Verify location captured\n5. Test already assigned message', '', ''),
        ('Quick Service', '1. Select Quick Service\n2. Scan asset\n3. Fill service form\n4. Submit\n5. Verify logged\n6. Verify status updated', '', ''),
        ('Quick Repair', '1. Select Quick Repair\n2. Scan asset\n3. Fill repair form\n4. Submit\n5. Verify logged\n6. Verify status updated', '', ''),
        ('Service Sign-off', '1. Open pending service task\n2. Fill sign-off form\n3. Attach report (optional)\n4. Sign off\n5. Verify status "In Service"\n6. Verify activity logged', '', ''),
        ('Repair Sign-off', '1. Open pending repair task\n2. Fill sign-off form\n3. Attach report (optional)\n4. Sign off\n5. Verify status "In Service"\n6. Verify activity logged', '', ''),
    ]
    
    for feature, steps, working, feedback in shortcuts_tests:
        row_cells = shortcuts_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 10: Admin Features ==========
    doc.add_heading('10. Admin Features', 1)
    
    admin_table = doc.add_table(rows=1, cols=4)
    admin_table.style = 'Light Grid Accent 1'
    
    hdr_cells = admin_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    admin_tests = [
        ('Admin Access', '1. Login as admin\n2. Verify admin features visible\n3. Verify Generate QR shortcut available', '', ''),
        ('Generate QR Sheet', '1. Click Generate QR Sheet\n2. Enter number of sheets\n3. Generate\n4. Verify PDF created\n5. Verify QR codes correct', '', ''),
        ('User Management', '1. Navigate to Admin\n2. View users list\n3. Verify user details\n4. Test user actions', '', ''),
        ('Domain Management', '1. Access domain management\n2. Verify domains listed\n3. Test domain operations', '', ''),
        ('Reset Password', '1. Select user\n2. Reset password\n3. Verify reset email sent', '', ''),
    ]
    
    for feature, steps, working, feedback in admin_tests:
        row_cells = admin_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 11: Profile & Settings ==========
    doc.add_heading('11. Profile & Settings', 1)
    
    profile_table = doc.add_table(rows=1, cols=4)
    profile_table.style = 'Light Grid Accent 1'
    
    hdr_cells = profile_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    profile_tests = [
        ('Profile View', '1. Navigate to Profile\n2. Verify user info displays\n3. Verify email/name shown', '', ''),
        ('Edit Profile', '1. Edit profile fields\n2. Save changes\n3. Verify updates saved', '', ''),
        ('My Assets', '1. Navigate to My Assets\n2. Verify assigned assets listed\n3. Verify can access assets', '', ''),
    ]
    
    for feature, steps, working, feedback in profile_tests:
        row_cells = profile_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 12: Mobile-Specific ==========
    doc.add_heading('12. Mobile-Specific Features', 1)
    
    mobile_table = doc.add_table(rows=1, cols=4)
    mobile_table.style = 'Light Grid Accent 1'
    
    hdr_cells = mobile_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    mobile_tests = [
        ('Bottom Tab Navigation', '1. Verify tabs visible\n2. Switch between tabs\n3. Verify navigation works', '', ''),
        ('Mobile Layout - Search', '1. Open search on mobile\n2. Verify card layout\n3. Verify responsive design\n4. Verify filters accessible', '', ''),
        ('Mobile Layout - Certs', '1. Open certs on mobile\n2. Verify card layout\n3. Verify alignment\n4. Verify filters work', '', ''),
        ('Screen Header', '1. Verify header displays\n2. Verify back button works\n3. Verify title centered\n4. Verify navigation consistent', '', ''),
        ('Touch Interactions', '1. Test all touch targets\n2. Verify adequate size\n3. Verify feedback on tap', '', ''),
        ('Keyboard Handling', '1. Open forms\n2. Verify keyboard appears\n3. Verify input accessible\n4. Verify keyboard dismisses', '', ''),
        ('Camera Permissions', '1. Request camera access\n2. Verify permission prompt\n3. Grant/deny\n4. Verify behavior', '', ''),
        ('Location Permissions', '1. Request location access\n2. Verify permission prompt\n3. Grant/deny\n4. Verify behavior', '', ''),
    ]
    
    for feature, steps, working, feedback in mobile_tests:
        row_cells = mobile_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 13: Web-Specific ==========
    doc.add_heading('13. Web-Specific Features', 1)
    
    web_table = doc.add_table(rows=1, cols=4)
    web_table.style = 'Light Grid Accent 1'
    
    hdr_cells = web_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    web_tests = [
        ('Web Navbar', '1. Verify navbar displays\n2. Verify all links work\n3. Verify active state highlighting\n4. Verify Certs link highlights', '', ''),
        ('Table Layout', '1. Verify table displays correctly\n2. Verify columns visible\n3. Verify horizontal scroll works\n4. Verify borders aligned', '', ''),
        ('Grid/Table Toggle', '1. Switch between views\n2. Verify both work\n3. Verify state persists', '', ''),
        ('Responsive Design', '1. Resize browser window\n2. Verify layout adapts\n3. Verify no horizontal scroll\n4. Verify all features accessible', '', ''),
        ('Keyboard Shortcuts', '1. Test keyboard navigation\n2. Verify shortcuts work\n3. Verify focus management', '', ''),
    ]
    
    for feature, steps, working, feedback in web_tests:
        row_cells = web_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== SECTION 14: Performance & Error Handling ==========
    doc.add_heading('14. Performance & Error Handling', 1)
    
    perf_table = doc.add_table(rows=1, cols=4)
    perf_table.style = 'Light Grid Accent 1'
    
    hdr_cells = perf_table.rows[0].cells
    hdr_cells[0].text = 'Feature/Input'
    hdr_cells[1].text = 'Test Steps'
    hdr_cells[2].text = 'Working'
    hdr_cells[3].text = 'Feedback/Issues'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    perf_tests = [
        ('Loading States', '1. Perform slow operations\n2. Verify loading indicators\n3. Verify user feedback', '', ''),
        ('Error Messages', '1. Trigger errors\n2. Verify error messages display\n3. Verify helpful messages\n4. Verify recovery options', '', ''),
        ('Network Errors', '1. Disconnect network\n2. Perform actions\n3. Verify error handling\n4. Verify retry options', '', ''),
        ('Large Data Sets', '1. Load large asset list\n2. Verify performance\n3. Verify pagination works\n4. Verify no crashes', '', ''),
        ('Image Loading', '1. Load assets with images\n2. Verify images load\n3. Verify placeholders\n4. Verify error handling', '', ''),
        ('Form Validation', '1. Submit invalid forms\n2. Verify validation messages\n3. Verify prevents submission\n4. Verify helpful hints', '', ''),
        ('Concurrent Actions', '1. Perform multiple actions\n2. Verify no conflicts\n3. Verify state consistency', '', ''),
    ]
    
    for feature, steps, working, feedback in perf_tests:
        row_cells = perf_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = steps
        row_cells[2].text = working
        row_cells[3].text = feedback
    
    doc.add_paragraph()
    
    # ========== Summary Section ==========
    doc.add_page_break()
    doc.add_heading('Test Summary', 1)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Total Test Cases: ').bold = True
    summary_para.add_run('Fill in total count after testing')
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Passed: ').bold = True
    summary_para.add_run('________')
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Failed: ').bold = True
    summary_para.add_run('________')
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Blocked: ').bold = True
    summary_para.add_run('________')
    
    doc.add_paragraph()
    doc.add_heading('Overall Feedback', 1)
    doc.add_paragraph('Use this section to provide overall feedback, major issues, and recommendations:')
    doc.add_paragraph()
    doc.add_paragraph('________________________________________________________________')
    doc.add_paragraph('________________________________________________________________')
    doc.add_paragraph('________________________________________________________________')
    
    # Save document
    output_path = 'Asset_Management_App_Test_Checklist.docx'
    doc.save(output_path)
    print(f'Test checklist generated: {output_path}')
    return output_path

if __name__ == '__main__':
    try:
        create_test_checklist()
    except ImportError:
        print('Error: python-docx library not found.')
        print('Install it with: pip install python-docx')
    except Exception as e:
        print(f'Error generating checklist: {e}')

